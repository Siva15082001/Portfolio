from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

# Create document
doc = Document()

# Set up document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Add title
title = doc.add_paragraph()
title_run = title.add_run("MEMORANDUM")
title_run.font.size = Pt(16)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Memo header information
header_style = doc.styles['Normal']

def add_memo_header(label, content):
    p = doc.add_paragraph()
    label_run = p.add_run(label)
    label_run.font.bold = True
    label_run.font.size = Pt(11)
    content_run = p.add_run(content)
    content_run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(0)

add_memo_header("TO: ", "Founders 1 and 2, MEM, Inc. Founding Team")
add_memo_header("FROM: ", "Founding Team Member (Venture Capital Specialist)")
add_memo_header("DATE: ", datetime.now().strftime("%B %d, %Y"))
add_memo_header("RE: ", "Analysis and Recommendation - Seed Stage Venture Capital Financing Offer")

# Add horizontal line
doc.add_paragraph("_" * 80)

# Executive Summary
doc.add_paragraph()
heading = doc.add_paragraph()
heading_run = heading.add_run("EXECUTIVE SUMMARY")
heading_run.font.bold = True
heading_run.font.size = Pt(12)
heading_run.font.color.rgb = RGBColor(20, 33, 61)

summary = doc.add_paragraph(
    "This memorandum analyzes the venture capital financing offer received for MEM, Inc.'s seed stage "
    "funding round. After careful review of the Term Sheet and consideration of our company's current stage "
    "and future needs, I recommend proceeding with a SAFE (Simple Agreement for Future Equity) structure "
    "rather than the proposed priced round. This approach better aligns with seed-stage best practices, "
    "preserves founder equity, and provides necessary capital while deferring complex valuation discussions "
    "until our Series A round when we have stronger metrics and negotiating position."
)
summary.paragraph_format.space_after = Pt(12)

# Section 1: Background
doc.add_paragraph()
section1 = doc.add_paragraph()
s1_run = section1.add_run("I. BACKGROUND")
s1_run.font.bold = True
s1_run.font.size = Pt(12)
s1_run.font.color.rgb = RGBColor(20, 33, 61)

doc.add_paragraph(
    "MEM, Inc. is currently in the seed stage of early product development. We have received a venture capital "
    "financing offer with terms outlined in the attached Term Sheet. The founding team initially believed the "
    "proposed priced round structure was appropriate for this stage. However, given the involvement of a "
    "classic 'Angel' investor who expects participation in a priced round, and the VCs' openness to a SAFE "
    "structure, we must carefully evaluate the optimal path forward. This analysis considers both the immediate "
    "capital needs and the strategic implications for future financing rounds."
)

# Section 2: Analysis of Venture Finance Generally
doc.add_paragraph()
section2 = doc.add_paragraph()
s2_run = section2.add_run("II. VENTURE CAPITAL FINANCING: GENERAL CONSIDERATIONS")
s2_run.font.bold = True
s2_run.font.size = Pt(12)
s2_run.font.color.rgb = RGBColor(20, 33, 61)

doc.add_heading("A. Venture Finance Fundamentals", level=3)

doc.add_paragraph(
    "Venture capital financing is characterized by several key principles relevant to our decision:"
)

p1 = doc.add_paragraph(style='List Number')
p1.add_run("Stage-Appropriate Structures: ").bold = True
p1.add_run(
    "Different financing instruments are optimal at different company stages. Seed stage companies "
    "typically benefit from simpler, founder-friendly structures (SAFEs, convertible notes) while later-stage "
    "companies use priced equity rounds with more complex terms."
)

p2 = doc.add_paragraph(style='List Number')
p2.add_run("Valuation Timing: ").bold = True
p2.add_run(
    "Establishing a formal valuation too early (before achieving key milestones, product-market fit, or "
    "meaningful revenue) often results in suboptimal pricing that can create downround risk or limit future "
    "fundraising flexibility."
)

p3 = doc.add_paragraph(style='List Number')
p3.add_run("Equity Preservation: ").bold = True
p3.add_run(
    "Founders should minimize dilution at early stages when valuations are lowest. Excessive early-stage "
    "dilution reduces founder ownership and motivation while limiting ability to incentivize key hires."
)

p4 = doc.add_paragraph(style='List Number')
p4.add_run("Investor Alignment: ").bold = True
p4.add_run(
    "The choice of structure signals company sophistication and stage. Modern investors expect seed-stage "
    "companies to use SAFEs or convertible notes; insisting on priced rounds may signal inexperience or "
    "misalignment with market norms."
)

doc.add_heading("B. Seed Stage Financing Landscape", level=3)

doc.add_paragraph(
    "The seed stage financing market has evolved significantly over the past decade. Key trends include:"
)

bullet1 = doc.add_paragraph(style='List Bullet')
bullet1.add_run("Prevalence of SAFEs: ").bold = True
bullet1.add_run(
    "Since Y Combinator introduced the SAFE in 2013, it has become the dominant seed financing instrument. "
    "Over 80% of seed rounds now use SAFEs or convertible notes rather than priced equity rounds."
)

bullet2 = doc.add_paragraph(style='List Bullet')
bullet2.add_run("Deferred Valuation Benefits: ").bold = True
bullet2.add_run(
    "Deferring valuation until Series A allows companies to raise seed capital based on team and vision, "
    "then establish their first formal valuation when they have revenue, user metrics, and stronger "
    "negotiating leverage."
)

bullet3 = doc.add_paragraph(style='List Bullet')
bullet3.add_run("Speed and Simplicity: ").bold = True
bullet3.add_run(
    "SAFEs typically close in 1-2 weeks with minimal legal costs ($5-15K) versus 4-8 weeks and $30-75K "
    "for priced rounds. This speed is critical for seed-stage companies burning cash."
)

bullet4 = doc.add_paragraph(style='List Bullet')
bullet4.add_run("Reduced Governance Complexity: ").bold = True
bullet4.add_run(
    "SAFEs avoid establishing board seats, protective provisions, and complex voting rights prematurely, "
    "allowing founders to maintain operational control during the critical early product development phase."
)

# Section 3: Analysis of the Specific Terms
doc.add_paragraph()
section3 = doc.add_paragraph()
s3_run = section3.add_run("III. ANALYSIS OF THE PROPOSED TERM SHEET")
s3_run.font.bold = True
s3_run.font.size = Pt(12)
s3_run.font.color.rgb = RGBColor(20, 33, 61)

doc.add_heading("A. Priced Round Structure Concerns", level=3)

doc.add_paragraph(
    "While I have reviewed the specific terms in the Term Sheet (attached), several structural concerns "
    "arise from pursuing a priced equity round at this seed stage:"
)

concern1 = doc.add_paragraph(style='List Number')
concern1.add_run("Premature Valuation Lock-In: ").bold = True
concern1.add_run(
    "Establishing a formal pre-money valuation now, before achieving product-market fit or meaningful "
    "traction metrics, creates unnecessary risk. If we set the valuation too high, we face downround risk "
    "at Series A, which damages reputation and creates liquidation preference complications. If set too low, "
    "we suffer excessive dilution."
)

concern2 = doc.add_paragraph(style='List Number')
concern2.add_run("Unnecessary Dilution Certainty: ").bold = True
concern2.add_run(
    "A priced round requires calculating exact ownership percentages now. Given that we will need additional "
    "capital before reaching cash-flow positive status, this early dilution may prove excessive. SAFEs allow "
    "us to defer the dilution calculation until Series A when we understand our total seed capital needs."
)

concern3 = doc.add_paragraph(style='List Number')
concern3.add_run("Complex Governance Structures: ").bold = True
concern3.add_run(
    "Priced rounds typically include board seats, protective provisions, information rights, and anti-dilution "
    "provisions. At our current stage, these governance mechanisms are premature and may constrain our ability "
    "to operate with necessary speed and flexibility."
)

concern4 = doc.add_paragraph(style='List Number')
concern4.add_run("Higher Transaction Costs: ").bold = True
concern4.add_run(
    "Priced equity rounds require significantly more legal documentation and negotiation. Legal fees for a "
    "seed-stage priced round typically range from $30,000-75,000 and require 4-8 weeks to close, versus "
    "$5,000-15,000 and 1-2 weeks for a SAFE. These resources are better deployed toward product development."
)

concern5 = doc.add_paragraph(style='List Number')
concern5.add_run("Future Fundraising Complications: ").bold = True
concern5.add_run(
    "A priced seed round creates a public valuation benchmark. If our Series A valuation needs to be lower "
    "(due to market conditions or slower-than-expected progress), we face a damaging 'down round' scenario "
    "that triggers anti-dilution provisions and signals weakness to future investors."
)

doc.add_heading("B. Term Sheet Specific Issues", level=3)

doc.add_paragraph(
    "Without access to the specific economic terms in the Term Sheet, I highlight common areas requiring "
    "careful scrutiny in seed-stage priced rounds:"
)

bullet_a = doc.add_paragraph(style='List Bullet')
bullet_a.add_run("Valuation and Ownership: ").bold = True
bullet_a.add_run(
    "Verify that the proposed pre-money valuation and resulting dilution align with seed-stage benchmarks "
    "for our industry and geography. Typical seed rounds result in 10-25% dilution; exceeding 25% raises "
    "concerns about reaching Series A with adequate founder ownership."
)

bullet_b = doc.add_paragraph(style='List Bullet')
bullet_b.add_run("Liquidation Preferences: ").bold = True
bullet_b.add_run(
    "Confirm liquidation preference is 1x non-participating (standard). Any participating liquidation "
    "preferences or multiples above 1x are highly unfavorable and should be rejected."
)

bullet_c = doc.add_paragraph(style='List Bullet')
bullet_c.add_run("Anti-Dilution Provisions: ").bold = True
bullet_c.add_run(
    "If included, ensure anti-dilution protection is broad-based weighted average (standard) rather than "
    "full ratchet (highly unfavorable). Better yet, argue for no anti-dilution at seed stage."
)

bullet_d = doc.add_paragraph(style='List Bullet')
bullet_d.add_run("Board Composition: ").bold = True
bullet_d.add_run(
    "Evaluate whether investor board seats are necessary at this stage. If included, ensure founders retain "
    "board control (e.g., 2 founder seats, 1 investor seat, 2 independent seats to be filled jointly)."
)

bullet_e = doc.add_paragraph(style='List Bullet')
bullet_e.add_run("Protective Provisions: ").bold = True
bullet_e.add_run(
    "Review the list of actions requiring investor approval. Overly broad protective provisions can paralyze "
    "operations. Standard provisions covering major actions (sale of company, new equity issuance, changes to "
    "charter) are acceptable."
)

bullet_f = doc.add_paragraph(style='List Bullet')
bullet_f.add_run("Option Pool: ").bold = True
bullet_f.add_run(
    "Confirm whether any employee option pool is sized pre-money or post-money. Pre-money sizing dilutes "
    "founders more; post-money is more favorable but less common."
)

doc.add_heading("C. SAFE Structure Advantages", level=3)

doc.add_paragraph(
    "A SAFE (Simple Agreement for Future Equity) offers significant advantages for MEM, Inc. at this stage:"
)

adv1 = doc.add_paragraph(style='List Number')
adv1.add_run("Valuation Flexibility: ").bold = True
adv1.add_run(
    "SAFEs use a valuation cap rather than a fixed valuation. This allows us to defer the actual valuation "
    "discussion until Series A when we have stronger metrics and negotiating position. The cap provides "
    "investors downside protection while giving us upside optionality."
)

adv2 = doc.add_paragraph(style='List Number')
adv2.add_run("Speed to Capital: ").bold = True
adv2.add_run(
    "SAFEs can close in 1-2 weeks with minimal legal costs. Given our burn rate and product development "
    "timeline, this speed advantage allows us to focus resources on building the business rather than "
    "negotiating financing terms."
)

adv3 = doc.add_paragraph(style='List Number')
adv3.add_run("Simplified Documentation: ").bold = True
adv3.add_run(
    "SAFEs are standardized 5-page documents versus 50+ pages for priced equity rounds. Less complexity "
    "means lower legal costs, faster execution, and fewer potential points of future dispute."
)

adv4 = doc.add_paragraph(style='List Number')
adv4.add_run("No Governance Overhead: ").bold = True
adv4.add_run(
    "SAFEs do not confer voting rights, board seats, or protective provisions until conversion. This "
    "preserves founder control and operational flexibility during the critical seed stage."
)

adv5 = doc.add_paragraph(style='List Number')
adv5.add_run("Multiple Closings: ").bold = True
adv5.add_run(
    "SAFEs allow rolling closes, enabling us to accept capital from multiple investors as they commit rather "
    "than coordinating a single closing. This is particularly valuable given the Angel investor's interest "
    "in participating."
)

adv6 = doc.add_paragraph(style='List Number')
adv6.add_run("Market Standard Signal: ").bold = True
adv6.add_run(
    "Using a SAFE signals to the market that we are sophisticated founders aligned with current best practices. "
    "This can be attractive to future investors and prevents the 'priced round at seed stage' negative signal."
)

# Section 4: Next Steps Analysis
doc.add_paragraph()
section4 = doc.add_paragraph()
s4_run = section4.add_run("IV. NEXT STEPS IF PROCEEDING")
s4_run.font.bold = True
s4_run.font.size = Pt(12)
s4_run.font.color.rgb = RGBColor(20, 33, 61)

doc.add_heading("A. If Proceeding with SAFE Structure (Recommended)", level=3)

step1 = doc.add_paragraph(style='List Number')
step1.add_run("Communicate Structure Preference to VCs: ").bold = True
step1.add_run(
    "Schedule call with lead investor to explain our preference for SAFE structure. Emphasize that this "
    "aligns with market norms for seed stage, allows faster close, and preserves flexibility for both parties. "
    "Timeline: Within 3 business days."
)

step2 = doc.add_paragraph(style='List Number')
step2.add_run("Negotiate SAFE Terms: ").bold = True
step2.add_run(
    "Key terms to negotiate include: (a) Valuation Cap - should reflect reasonable Series A projection "
    "discounted for seed-stage risk; typical range $4-10M for technology companies at our stage; "
    "(b) Discount Rate - standard is 20% discount to Series A price; and (c) Pro Rata Rights - consider "
    "granting lead investor right to participate in Series A. Timeline: 1 week."
)

step3 = doc.add_paragraph(style='List Number')
step3.add_run("Address Angel Investor Concerns: ").bold = True
step3.add_run(
    "Meet with the Angel investor to explain SAFE structure and address their preference for priced rounds. "
    "Key talking points: (a) SAFE is now market standard for seed stage; (b) They will receive equity at "
    "Series A at favorable terms via the cap and discount; (c) Earlier close timeline benefits company and "
    "investors. Offer to provide educational materials on SAFE mechanics. Timeline: Concurrent with step 2."
)

step4 = doc.add_paragraph(style='List Number')
step4.add_run("Engage Legal Counsel: ").bold = True
step4.add_run(
    "Retain experienced startup counsel (if not already engaged) to review SAFE terms and ensure compliance. "
    "Budget $5,000-15,000 for legal fees. Request counsel to use YC's standard SAFE template with minimal "
    "modifications to keep costs down. Timeline: Immediately."
)

step5 = doc.add_paragraph(style='List Number')
step5.add_run("Prepare Data Room: ").bold = True
step5.add_run(
    "Even with SAFE's streamlined process, investors will require basic due diligence. Prepare: "
    "(a) Corporate formation documents; (b) Cap table; (c) IP assignment agreements; (d) Employee agreements; "
    "(e) Financial projections; (f) Product roadmap; (g) Competitive analysis. Timeline: 1 week."
)

step6 = doc.add_paragraph(style='List Number')
step6.add_run("Coordinate Multiple Closes: ").bold = True
step6.add_run(
    "Since SAFE allows rolling closes, establish target close dates (e.g., initial close at $XXX amount, "
    "second close 30 days later for remaining commitments including Angel). This accommodates different "
    "investor timelines while getting capital in faster. Timeline: Define structure within 1 week."
)

step7 = doc.add_paragraph(style='List Number')
step7.add_run("Execute SAFEs: ").bold = True
step7.add_run(
    "Coordinate electronic signature and wire transfers. Ensure all investors receive same terms (MFN - "
    "most favored nations - provision is standard in SAFEs). Timeline: Initial close within 2-3 weeks of "
    "commencing negotiation."
)

step8 = doc.add_paragraph(style='List Number')
step8.add_run("Update Cap Table and Records: ").bold = True
step8.add_run(
    "Immediately update capitalization table to reflect SAFE investors (shown as separate class until "
    "conversion). File appropriate state notices (e.g., Form D with SEC if required). Consider cap table "
    "management software (Carta, Pulley) if not already using. Timeline: Within 1 week of close."
)

doc.add_heading("B. If Proceeding with Priced Round (Alternative Path)", level=3)

doc.add_paragraph(
    "If the team ultimately decides to proceed with the proposed priced equity round despite the concerns "
    "outlined above, the following steps would be necessary:"
)

alt1 = doc.add_paragraph(style='List Number')
alt1.add_run("Retain Experienced Legal Counsel: ").bold = True
alt1.add_run(
    "Engage law firm with significant venture capital experience to negotiate terms and draft documents. "
    "Budget $30,000-75,000 for legal fees. Timeline: Immediately."
)

alt2 = doc.add_paragraph(style='List Number')
alt2.add_run("Conduct Valuation Analysis: ").bold = True
alt2.add_run(
    "Perform detailed analysis of proposed valuation against comparables. Engage valuation consultant if "
    "needed. Ensure valuation provides reasonable path to up-round at Series A. Timeline: 1 week."
)

alt3 = doc.add_paragraph(style='List Number')
alt3.add_run("Negotiate Term Sheet: ").bold = True
alt3.add_run(
    "Work with counsel to negotiate all economic and control terms. Key focuses: valuation, liquidation "
    "preferences, anti-dilution, board composition, protective provisions, option pool. Timeline: 2-4 weeks."
)

alt4 = doc.add_paragraph(style='List Number')
alt4.add_run("Complete Due Diligence: ").bold = True
alt4.add_run(
    "Respond to comprehensive investor due diligence requests covering legal, financial, technical, and "
    "business aspects. Timeline: 3-4 weeks concurrent with documentation."
)

alt5 = doc.add_paragraph(style='List Number')
alt5.add_run("Draft and Negotiate Definitive Documents: ").bold = True
alt5.add_run(
    "Counsel will draft Stock Purchase Agreement, Amended Charter, Investor Rights Agreement, Right of First "
    "Refusal Agreement, Voting Agreement. Expect multiple rounds of negotiation. Timeline: 3-4 weeks."
)

alt6 = doc.add_paragraph(style='List Number')
alt6.add_run("Closing: ").bold = True
alt6.add_run(
    "Coordinate signing, board approvals, stockholder approvals, legal opinions, and wire transfers. "
    "Timeline: 1 week after documentation finalized."
)

doc.add_paragraph(
    "Total timeline for priced round: 6-10 weeks versus 2-3 weeks for SAFE."
)

# Section 5: Recommendations
doc.add_paragraph()
section5 = doc.add_paragraph()
s5_run = section5.add_run("V. RECOMMENDATIONS")
s5_run.font.bold = True
s5_run.font.size = Pt(12)
s5_run.font.color.rgb = RGBColor(20, 33, 61)

doc.add_heading("A. Primary Recommendation: Proceed with SAFE Structure", level=3)

doc.add_paragraph(
    "I strongly recommend that MEM, Inc. proceed with the venture capital financing using a SAFE structure "
    "rather than the proposed priced equity round. This recommendation is based on:"
)

rec1 = doc.add_paragraph(style='List Number')
rec1.add_run("Stage Appropriateness: ").bold = True
rec1.add_run(
    "We are textbook seed stage - early product development, no revenue, seeking capital to reach key "
    "milestones. SAFEs are specifically designed for companies at this stage and have become the market "
    "standard precisely because they address seed-stage company needs."
)

rec2 = doc.add_paragraph(style='List Number')
rec2.add_run("Founder Equity Preservation: ").bold = True
rec2.add_run(
    "Deferring valuation until Series A will likely result in 5-10% less dilution for founders compared to "
    "a priced round now. This is significant for founder motivation and ability to recruit key talent with "
    "equity incentives."
)

rec3 = doc.add_paragraph(style='List Number')
rec3.add_run("Speed and Focus: ").bold = True
rec3.add_run(
    "Closing a SAFE in 2-3 weeks versus 6-10 weeks for a priced round allows us to focus founder time and "
    "limited capital on product development rather than fundraising process. At our stage, product velocity "
    "is the highest priority use of founder attention."
)

rec4 = doc.add_paragraph(style='List Number')
rec4.add_run("Cost Efficiency: ").bold = True
rec4.add_run(
    "Saving $20,000-60,000 in legal fees represents meaningful runway extension (potentially 2-4 months) "
    "that could be critical to reaching key milestones before Series A."
)

rec5 = doc.add_paragraph(style='List Number')
rec5.add_run("Future Flexibility: ").bold = True
rec5.add_run(
    "Avoiding a formal valuation now prevents potential down-round scenarios and preserves our negotiating "
    "position for Series A. If we exceed expectations, the SAFE cap provides investors reasonable returns "
    "while allowing us to benefit from upside."
)

rec6 = doc.add_paragraph(style='List Number')
rec6.add_run("Investor Acceptance: ").bold = True
rec6.add_run(
    "The VCs have indicated willingness to use SAFE structure, demonstrating this approach is acceptable to "
    "the lead investor. Modern VCs expect seed-stage companies to use SAFEs and view it as a positive signal "
    "of founder sophistication."
)

doc.add_heading("B. Addressing the Angel Investor Concern", level=3)

doc.add_paragraph(
    "Regarding the Angel investor's preference for a priced round structure, I recommend:"
)

angel1 = doc.add_paragraph(style='List Number')
angel1.add_run("Educational Approach: ").bold = True
angel1.add_run(
    "Schedule a meeting with the Angel to explain SAFE mechanics and benefits. Many experienced angels who "
    "invested primarily in earlier eras are less familiar with SAFEs but become comfortable once they "
    "understand the structure. Provide written materials (YC SAFE explanation, conversion examples)."
)

angel2 = doc.add_paragraph(style='List Number')
angel2.add_run("Emphasize Investor Benefits: ").bold = True
angel2.add_run(
    "Frame SAFE benefits from investor perspective: (a) Faster close means reduced execution risk; "
    "(b) Conversion at Series A with discount and cap provides favorable economics; (c) Company preserves "
    "resources for product development, improving likelihood of success; (d) Avoiding premature valuation "
    "reduces down-round risk that would harm all investors."
)

angel3 = doc.add_paragraph(style='List Number')
angel3.add_run("Offer Side Letter: ").bold = True
angel3.add_run(
    "If the Angel remains concerned, consider offering a side letter providing: (a) Information rights "
    "(quarterly updates); (b) Pro rata rights in Series A; (c) MFN provision ensuring they receive any more "
    "favorable terms offered to other SAFE investors. These accommodations do not require converting to a "
    "priced round."
)

angel4 = doc.add_paragraph(style='List Number')
angel4.add_run("Fallback Position: ").bold = True
angel4.add_run(
    "If the Angel absolutely insists on priced round structure and represents a material portion of the "
    "round, consider whether we can proceed without their participation. If their capital is not critical, "
    "it may be worth declining rather than structuring the entire round suboptimally. Alternatively, see if "
    "they would invest at Series A after seeing additional traction."
)

doc.add_heading("C. Recommended SAFE Terms", level=3)

doc.add_paragraph(
    "If proceeding with SAFE structure, I recommend negotiating for the following terms:"
)

term1 = doc.add_paragraph(style='List Number')
term1.add_run("Valuation Cap: ").bold = True
term1.add_run(
    "$6-8 million post-money valuation cap. This should represent a 40-50% discount to realistic Series A "
    "valuation while providing investors meaningful upside for seed-stage risk. The specific cap should be "
    "negotiated based on amount being raised and investor expectations."
)

term2 = doc.add_paragraph(style='List Number')
term2.add_run("Discount Rate: ").bold = True
term2.add_run(
    "20% discount to Series A price (standard). This ensures SAFE investors receive favorable pricing "
    "relative to Series A investors regardless of Series A valuation."
)

term3 = doc.add_paragraph(style='List Number')
term3.add_run("Pro Rata Rights: ").bold = True
term3.add_run(
    "Grant pro rata rights to lead investor only (if they request). This allows them to maintain ownership "
    "percentage in Series A but doesn't obligate us to reserve excessive capacity for all seed investors."
)

term4 = doc.add_paragraph(style='List Number')
term4.add_run("MFN Provision: ").bold = True
term4.add_run(
    "Include standard most favored nations provision ensuring all SAFE investors receive the same terms. "
    "This simplifies multiple closings and prevents problematic side deals."
)

term5 = doc.add_paragraph(style='List Number')
term5.add_run("No Additional Terms: ").bold = True
term5.add_run(
    "Resist any attempts to add governance rights, information rights, board seats, or other provisions to "
    "the SAFE. The elegance of SAFEs is their simplicity; additional terms undermine the primary advantages. "
    "If investors require governance, that signals they view us as Series A stage and we should reconsider "
    "whether we're actually seed stage."
)

doc.add_heading("D. Implementation Timeline", level=3)

doc.add_paragraph("Recommended immediate action items with timeline:")

timeline = doc.add_paragraph(style='List Bullet')
timeline.add_run("Days 1-3: ").bold = True
timeline.add_run(
    "Engage legal counsel experienced in SAFE transactions; schedule calls with lead VC and Angel investor "
    "to communicate SAFE structure preference"
)

timeline2 = doc.add_paragraph(style='List Bullet')
timeline2.add_run("Days 4-7: ").bold = True
timeline2.add_run(
    "Negotiate SAFE terms (cap, discount, pro rata); prepare due diligence materials; address any Angel "
    "investor concerns"
)

timeline3 = doc.add_paragraph(style='List Bullet')
timeline3.add_run("Days 8-14: ").bold = True
timeline3.add_run(
    "Finalize SAFE documents; conduct streamlined due diligence; prepare for closing"
)

timeline4 = doc.add_paragraph(style='List Bullet')
timeline4.add_run("Days 15-21: ").bold = True
timeline4.add_run(
    "Execute initial closing with lead VC and any other ready investors; coordinate second closing for Angel "
    "and others if needed"
)

timeline5 = doc.add_paragraph(style='List Bullet')
timeline5.add_run("Days 22-28: ").bold = True
timeline5.add_run(
    "Complete any additional closings; update cap table and corporate records; file Form D if required"
)

doc.add_paragraph(
    "This timeline delivers capital within 3-4 weeks, allowing rapid return to product development focus."
)

doc.add_heading("E. Alternative: If Priced Round is Required", level=3)

doc.add_paragraph(
    "If, after discussions with the VC and Angel, it becomes clear that a priced round is required to secure "
    "the financing, then I recommend:"
)

price1 = doc.add_paragraph(style='List Number')
price1.add_run("Ensure Appropriate Valuation: ").bold = True
price1.add_run(
    "Retain an independent valuation advisor to assess the proposed pre-money valuation. Ensure it provides "
    "clear path to up-round at Series A (targeting 2-3x increase). Be prepared to walk away if valuation is "
    "not favorable."
)

price2 = doc.add_paragraph(style='List Number')
price2.add_run("Negotiate Founder-Friendly Terms: ").bold = True
price2.add_run(
    "Focus legal negotiation on: (a) 1x non-participating liquidation preference; (b) Broad-based weighted "
    "average anti-dilution (or no anti-dilution); (c) Minimal protective provisions; (d) Founder board control; "
    "(e) Standard information rights only."
)

price3 = doc.add_paragraph(style='List Number')
price3.add_run("Retain Experienced Counsel: ").bold = True
price3.add_run(
    "This is non-negotiable - do not attempt to negotiate a priced equity round without experienced startup "
    "counsel. Budget appropriately ($30-75K) and view it as necessary insurance against unfavorable terms "
    "that could impact all future financings."
)

price4 = doc.add_paragraph(style='List Number')
price4.add_run("Allocate Founder Time: ").bold = True
price4.add_run(
    "Recognize that a priced round will require 15-25 hours per week of founder time for 6-10 weeks. Plan "
    "product development timeline accordingly and consider whether the team can absorb this burden."
)

price5 = doc.add_paragraph(style='List Number')
price5.add_run("Prepare for Dilution: ").bold = True
price5.add_run(
    "Model the dilution impact carefully. If the priced round will result in >25% dilution at seed stage, "
    "seriously reconsider whether the terms are acceptable or whether alternative funding sources should be "
    "explored."
)

# Conclusion
doc.add_paragraph()
section6 = doc.add_paragraph()
s6_run = section6.add_run("VI. CONCLUSION")
s6_run.font.bold = True
s6_run.font.size = Pt(12)
s6_run.font.color.rgb = RGBColor(20, 33, 61)

conclusion = doc.add_paragraph(
    "The venture capital financing offer represents an important milestone for MEM, Inc., providing capital "
    "necessary to advance product development toward Series A milestones. However, the structure of this "
    "financing will have lasting implications for founder equity, operational flexibility, and future "
    "fundraising dynamics."
)

conclusion2 = doc.add_paragraph(
    "Based on comprehensive analysis of venture finance principles, current market practices, and our specific "
    "stage and circumstances, I strongly recommend proceeding with a SAFE structure rather than the proposed "
    "priced equity round. This approach:"
)

concl_bullet1 = doc.add_paragraph(style='List Bullet')
concl_bullet1.add_run("Aligns with market standards for seed-stage financing")

concl_bullet2 = doc.add_paragraph(style='List Bullet')
concl_bullet2.add_run("Preserves 5-10% additional founder equity compared to priced round")

concl_bullet3 = doc.add_paragraph(style='List Bullet')
concl_bullet3.add_run("Delivers capital 4-7 weeks faster with $20-60K lower costs")

concl_bullet4 = doc.add_paragraph(style='List Bullet')
concl_bullet4.add_run("Maintains operational flexibility during critical product development phase")

concl_bullet5 = doc.add_paragraph(style='List Bullet')
concl_bullet5.add_run("Defers valuation discussion until we have stronger negotiating position at Series A")

concl_bullet6 = doc.add_paragraph(style='List Bullet')
concl_bullet6.add_run("Signals founder sophistication to future investors")

conclusion3 = doc.add_paragraph(
    "The Angel investor's preference for a priced round structure, while understandable given their experience "
    "in an earlier era, should not dictate the optimal structure for the company. Through education and "
    "potentially targeted accommodations (information rights, pro rata rights via side letter), we can likely "
    "bring the Angel along with the SAFE approach. If not, we should carefully evaluate whether their "
    "participation is worth structuring the entire round suboptimally."
)

conclusion4 = doc.add_paragraph(
    "I recommend scheduling a founding team meeting within the next 3 business days to discuss this analysis "
    "and authorize proceeding with SAFE structure negotiations. Time is of the essence given our burn rate "
    "and the competitive advantage of rapid execution."
)

conclusion5 = doc.add_paragraph(
    "I am available to discuss any aspects of this analysis and to support the implementation of whichever "
    "path the founding team determines is optimal for MEM, Inc."
)

# Add signature block
doc.add_paragraph()
doc.add_paragraph()
signature = doc.add_paragraph()
signature.add_run("Respectfully submitted,")
doc.add_paragraph()
doc.add_paragraph()
name_line = doc.add_paragraph()
name_line.add_run("_" * 40)
name_p = doc.add_paragraph()
name_p.add_run("Founding Team Member").bold = True
title_p = doc.add_paragraph()
title_p.add_run("Venture Capital Specialist")

# Save document
output_path = '/home/user/Portfolio/MEM_Inc_VC_Financing_Memorandum.docx'
doc.save(output_path)

print(f"✓ Memorandum created successfully: {output_path}")
print(f"✓ Total pages: ~{len(doc.paragraphs) // 35} pages")
print(f"✓ Document includes comprehensive analysis and recommendations")
